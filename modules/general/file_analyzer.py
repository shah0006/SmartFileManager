"""
General File Analysis Module.

This module provides functionality for analyzing and categorizing
general files that don't fit into specialized domains.
"""

import os
import re
import logging
import mimetypes
import datetime
from typing import Dict, List, Any, Optional, Tuple, Set
from pathlib import Path
import hashlib

from core.metadata_extractor import LLMMetadataExtractor

# Setup logging
logger = logging.getLogger(__name__)

class FileAnalyzer:
    """General file analysis functionality."""
    
    def __init__(self, llm_handler=None):
        """
        Initialize the file analyzer.
        
        Args:
            llm_handler: Handler for local LLM interactions
        """
        self.llm_handler = llm_handler
        
        # Initialize mimetypes
        mimetypes.init()
        
        # File categories mapping
        self.category_mapping = {
            # Documents
            'document': ['.pdf', '.doc', '.docx', '.rtf', '.txt', '.odt', '.md', '.tex'],
            
            # Spreadsheets
            'spreadsheet': ['.xls', '.xlsx', '.csv', '.ods', '.numbers'],
            
            # Presentations
            'presentation': ['.ppt', '.pptx', '.key', '.odp'],
            
            # Images
            'image': ['.jpg', '.jpeg', '.png', '.gif', '.bmp', '.tiff', '.tif', '.svg', '.webp', '.heic'],
            
            # Audio
            'audio': ['.mp3', '.wav', '.ogg', '.flac', '.aac', '.m4a', '.wma'],
            
            # Video
            'video': ['.mp4', '.avi', '.mov', '.wmv', '.flv', '.mkv', '.webm', '.m4v', '.mpg', '.mpeg'],
            
            # Archives
            'archive': ['.zip', '.rar', '.7z', '.tar', '.gz', '.bz2', '.xz', '.tgz'],
            
            # Code
            'code': ['.py', '.java', '.js', '.html', '.css', '.c', '.cpp', '.h', '.php', '.rb', '.go', '.ts', '.json', '.xml', '.sql', '.sh', '.bat', '.ps1'],
            
            # Executables
            'executable': ['.exe', '.app', '.msi', '.dmg', '.deb', '.rpm'],
            
            # Fonts
            'font': ['.ttf', '.otf', '.woff', '.woff2', '.eot'],
            
            # System
            'system': ['.dll', '.sys', '.config', '.ini'],
            
            # Data
            'data': ['.db', '.sqlite', '.sql', '.json', '.xml', '.yaml', '.yml', '.toml']
        }
        
        # Build reverse lookup for efficient categorization
        self.extension_to_category = {}
        for category, extensions in self.category_mapping.items():
            for ext in extensions:
                self.extension_to_category[ext] = category
    
    def categorize_file(self, file_path: str) -> Tuple[str, Dict[str, Any]]:
        """
        Categorize a file based on its extension and content.
        
        Args:
            file_path: Path to the file
            
        Returns:
            Tuple of (category, metadata)
        """
        metadata = self.extract_metadata(file_path)
        
        # Get file extension (lowercase)
        ext = os.path.splitext(file_path)[1].lower()
        
        # Check if extension is in our mapping
        if ext in self.extension_to_category:
            category = self.extension_to_category[ext]
        else:
            # If extension not found, try to determine from mimetype
            mimetype, _ = mimetypes.guess_type(file_path)
            
            if mimetype:
                if mimetype.startswith('image/'):
                    category = 'image'
                elif mimetype.startswith('audio/'):
                    category = 'audio'
                elif mimetype.startswith('video/'):
                    category = 'video'
                elif mimetype.startswith('text/'):
                    category = 'document'
                elif 'zip' in mimetype or 'compressed' in mimetype:
                    category = 'archive'
                elif 'executable' in mimetype:
                    category = 'executable'
                else:
                    category = 'other'
            else:
                category = 'other'
        
        return category, metadata
    
    def extract_metadata(self, file_path: str) -> Dict[str, Any]:
        """
        Extract basic metadata from a general file.
        
        Args:
            file_path: Path to the file
            
        Returns:
            Dictionary of metadata
        """
        metadata = {}
        
        try:
            # Get file stats
            stat = os.stat(file_path)
            
            # Basic file metadata
            metadata['filename'] = os.path.basename(file_path)
            metadata['extension'] = os.path.splitext(file_path)[1].lower()
            metadata['size'] = stat.st_size
            metadata['created'] = datetime.datetime.fromtimestamp(stat.st_ctime).isoformat()
            metadata['modified'] = datetime.datetime.fromtimestamp(stat.st_mtime).isoformat()
            metadata['accessed'] = datetime.datetime.fromtimestamp(stat.st_atime).isoformat()
            
            # Calculate hash for file identification
            metadata['md5'] = self._calculate_file_hash(file_path, algorithm='md5')
            
            # Try to extract content-based metadata for certain file types
            if self._should_extract_content_metadata(file_path):
                content_metadata = self._extract_content_metadata(file_path)
                metadata.update(content_metadata)
                
        except Exception as e:
            logger.error(f"Error extracting metadata from {file_path}: {e}")
        
        return metadata
    
    def _should_extract_content_metadata(self, file_path: str) -> bool:
        """
        Determine if content metadata should be extracted for this file.
        
        Args:
            file_path: Path to the file
            
        Returns:
            Whether to extract content metadata
        """
        # Get file size
        try:
            size = os.path.getsize(file_path)
            
            # Skip large files (>50 MB)
            if size > 50 * 1024 * 1024:
                return False
                
            # Get extension
            ext = os.path.splitext(file_path)[1].lower()
            
            # List of extensions for which we should extract content metadata
            content_extensions = set([
                '.txt', '.md', '.csv', '.json', '.xml', '.html', '.htm',
                '.pdf', '.doc', '.docx', '.rtf', '.odt'
            ])
            
            return ext in content_extensions
            
        except Exception as e:
            logger.error(f"Error checking file for content extraction: {e}")
            return False
    
    def _extract_content_metadata(self, file_path: str) -> Dict[str, Any]:
        """
        Extract metadata from file content.
        
        Args:
            file_path: Path to the file
            
        Returns:
            Dictionary of content metadata
        """
        metadata = {}
        
        try:
            ext = os.path.splitext(file_path)[1].lower()
            
            # Text files
            if ext in ['.txt', '.md', '.csv', '.json', '.xml', '.html', '.htm']:
                try:
                    with open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
                        content = f.read(4096)  # Read first 4KB for analysis
                        
                        # Count lines
                        metadata['line_count'] = content.count('\n') + 1
                        
                        # Estimate word count
                        words = re.findall(r'\w+', content)
                        metadata['word_count'] = len(words)
                        
                        # Extract language if LLM is available
                        if self.llm_handler and len(content.strip()) > 100:
                            metadata['language'] = self._detect_language(content[:1000])
                except Exception as e:
                    logger.error(f"Error extracting text metadata: {e}")
            
            # Document files
            elif ext in ['.pdf', '.doc', '.docx', '.rtf', '.odt']:
                # Use appropriate library based on file type
                if ext == '.pdf':
                    metadata.update(self._extract_pdf_metadata(file_path))
                elif ext in ['.doc', '.docx']:
                    metadata.update(self._extract_msword_metadata(file_path))
                
        except Exception as e:
            logger.error(f"Error extracting content metadata: {e}")
        
        return metadata
    
    def _extract_pdf_metadata(self, file_path: str) -> Dict[str, Any]:
        """Extract metadata from PDF file."""
        metadata = {}
        
        try:
            # Try PyMuPDF (fitz)
            try:
                import fitz
                
                doc = fitz.open(file_path)
                
                # Get document info
                info = doc.metadata
                if info:
                    if info.get('title'):
                        metadata['title'] = info['title']
                    if info.get('author'):
                        metadata['author'] = info['author']
                    if info.get('subject'):
                        metadata['subject'] = info['subject']
                    if info.get('keywords'):
                        metadata['keywords'] = info['keywords']
                
                # Get page count
                metadata['page_count'] = doc.page_count
                
                # Get text from first page for content analysis
                if doc.page_count > 0:
                    first_page = doc[0]
                    text = first_page.get_text()
                    metadata['first_page_text'] = text[:1000] if len(text) > 1000 else text
                
                doc.close()
                
            except ImportError:
                logger.warning("PyMuPDF not available for PDF metadata extraction")
                
        except Exception as e:
            logger.error(f"Error extracting PDF metadata: {e}")
        
        return metadata
    
    def _extract_msword_metadata(self, file_path: str) -> Dict[str, Any]:
        """Extract metadata from Microsoft Word file."""
        metadata = {}
        
        try:
            # Try python-docx for docx files
            if file_path.lower().endswith('.docx'):
                try:
                    import docx
                    
                    doc = docx.Document(file_path)
                    
                    # Get core properties
                    core_props = doc.core_properties
                    if core_props:
                        if core_props.title:
                            metadata['title'] = core_props.title
                        if core_props.author:
                            metadata['author'] = core_props.author
                        if core_props.subject:
                            metadata['subject'] = core_props.subject
                        if core_props.keywords:
                            metadata['keywords'] = core_props.keywords
                    
                    # Count paragraphs
                    metadata['paragraph_count'] = len(doc.paragraphs)
                    
                    # Extract some text for content analysis
                    text_parts = []
                    for i, para in enumerate(doc.paragraphs):
                        if i >= 10:  # Limit to first 10 paragraphs
                            break
                        if para.text:
                            text_parts.append(para.text)
                    
                    if text_parts:
                        metadata['sample_text'] = ' '.join(text_parts)[:1000]
                    
                except ImportError:
                    logger.warning("python-docx not available for Word metadata extraction")
            
        except Exception as e:
            logger.error(f"Error extracting Word metadata: {e}")
        
        return metadata
    
    def _detect_language(self, text: str) -> str:
        """
        Detect language of text using LLM.
        
        Args:
            text: Text sample
            
        Returns:
            Detected language
        """
        if not self.llm_handler:
            return "unknown"
        
        try:
            prompt = f"""
            Identify the language of the following text. Respond with just the language name in English, e.g., "English", "Spanish", "Chinese", etc.
            
            Text: {text[:500]}
            """
            
            response = self.llm_handler.get_response(prompt).strip()
            return response
            
        except Exception as e:
            logger.error(f"Error detecting language: {e}")
            return "unknown"
    
    def _calculate_file_hash(self, file_path: str, algorithm: str = 'md5', chunk_size: int = 8192) -> str:
        """
        Calculate hash of a file for identification.
        
        Args:
            file_path: Path to the file
            algorithm: Hash algorithm to use ('md5', 'sha1', 'sha256')
            chunk_size: Size of chunks to read
            
        Returns:
            Hexadecimal hash string
        """
        try:
            if algorithm == 'md5':
                hash_obj = hashlib.md5()
            elif algorithm == 'sha1':
                hash_obj = hashlib.sha1()
            elif algorithm == 'sha256':
                hash_obj = hashlib.sha256()
            else:
                raise ValueError(f"Unsupported hash algorithm: {algorithm}")
            
            with open(file_path, 'rb') as f:
                for chunk in iter(lambda: f.read(chunk_size), b''):
                    hash_obj.update(chunk)
            
            return hash_obj.hexdigest()
            
        except Exception as e:
            logger.error(f"Error calculating file hash: {e}")
            return ""
    
    def identify_duplicates(self, files: List[str]) -> Dict[str, List[str]]:
        """
        Identify duplicate files based on content hash.
        
        Args:
            files: List of file paths to check
            
        Returns:
            Dictionary mapping hash to lists of duplicate files
        """
        hash_map = {}
        
        for file_path in files:
            try:
                file_hash = self._calculate_file_hash(file_path)
                if file_hash:
                    if file_hash not in hash_map:
                        hash_map[file_hash] = []
                    hash_map[file_hash].append(file_path)
            except Exception as e:
                logger.error(f"Error processing file {file_path}: {e}")
        
        # Filter to only include hashes with multiple files (duplicates)
        duplicates = {h: files for h, files in hash_map.items() if len(files) > 1}
        
        return duplicates
